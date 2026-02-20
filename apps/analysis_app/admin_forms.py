from django import forms
from docx import Document

from apps.analysis_app.forms import GENERAL_SUMMARY_PU_LABEL
from apps.analysis_app.models import PortalDbConnectionSettings, SvodkaTemplate
from apps.analysis_app.pu_detection import detect_pu_from_docx
from apps.portaldb.gateway import get_portal_gateway


AUTO_PU_CHOICE_VALUE = "__AUTO__"
GENERAL_PU_CHOICE_VALUE = "__GENERAL__"


class PortalDbConnectionSettingsAdminForm(forms.ModelForm):
    password = forms.CharField(
        required=False,
        label="Пароль",
        widget=forms.PasswordInput(render_value=False, attrs={"placeholder": "********"}),
    )

    class Meta:
        model = PortalDbConnectionSettings
        fields = ("profile", "host", "port", "db_name", "user", "password")

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        if self.instance and self.instance.password_encrypted:
            self.fields["password"].help_text = (
                "Пароль сохранён. Оставьте пустым, чтобы не менять."
            )

    def save(self, commit=True):
        obj = super().save(commit=False)
        if commit:
            obj.save()
        return obj


class SvodkaTemplateAdminForm(forms.ModelForm):
    pu_select = forms.ChoiceField(label="Пограничное управление", required=False)

    class Meta:
        model = SvodkaTemplate
        fields = ("pu_select", "pu_name", "file", "begin_marker", "end_marker", "is_active")

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._pu_name_map: dict[str, str] = {}
        self.fields["pu_select"].choices = self._build_pu_choices()

        if self.instance and self.instance.pk:
            if self.instance.scope == SvodkaTemplate.Scope.GENERAL:
                self.fields["pu_select"].initial = GENERAL_PU_CHOICE_VALUE
            elif self.instance.pu_id:
                self.fields["pu_select"].initial = str(self.instance.pu_id)
            else:
                self.fields["pu_select"].initial = AUTO_PU_CHOICE_VALUE
        else:
            self.fields["pu_select"].initial = AUTO_PU_CHOICE_VALUE

    def _build_pu_choices(self):
        choices = [
            (AUTO_PU_CHOICE_VALUE, "Авто (определить по шаблону)"),
            (GENERAL_PU_CHOICE_VALUE, GENERAL_SUMMARY_PU_LABEL),
        ]
        try:
            gateway = get_portal_gateway()
            for pu in gateway.list_pus():
                pu_id = str(pu.pu_id)
                pu_name = str(pu.full_name or pu.short_name or "")
                self._pu_name_map[pu_id] = pu_name
                choices.append((pu_id, pu_name))
        except Exception:  # noqa: BLE001
            pass
        return choices

    def _open_docx_document(self, file_obj):
        if not file_obj:
            return None
        if hasattr(file_obj, "temporary_file_path"):
            return Document(file_obj.temporary_file_path())
        if hasattr(file_obj, "path"):
            return Document(file_obj.path)
        if hasattr(file_obj, "seek"):
            file_obj.seek(0)
        return Document(file_obj)

    def clean(self):
        cleaned_data = super().clean()
        pu_select = cleaned_data.get("pu_select")
        manual_pu_name = str(cleaned_data.get("pu_name") or "").strip()
        is_manual_override = "pu_name" in self.changed_data and bool(manual_pu_name)

        resolved_scope = SvodkaTemplate.Scope.GENERAL
        resolved_pu_id = ""
        resolved_pu_name = GENERAL_SUMMARY_PU_LABEL

        if pu_select == AUTO_PU_CHOICE_VALUE:
            file_obj = cleaned_data.get("file") or getattr(self.instance, "file", None)
            try:
                document = self._open_docx_document(file_obj)
                detection = detect_pu_from_docx(document) if document else None
            except Exception:  # noqa: BLE001
                detection = None
            if detection and detection.pu:
                resolved_scope = SvodkaTemplate.Scope.PU
                resolved_pu_id = str(detection.pu.portal_pu_id)
                resolved_pu_name = str(detection.pu.full_name or detection.pu.short_name or "")
        elif pu_select == GENERAL_PU_CHOICE_VALUE or not pu_select:
            resolved_scope = SvodkaTemplate.Scope.GENERAL
            resolved_pu_id = ""
            resolved_pu_name = GENERAL_SUMMARY_PU_LABEL
        else:
            resolved_scope = SvodkaTemplate.Scope.PU
            resolved_pu_id = str(pu_select)
            resolved_pu_name = self._pu_name_map.get(resolved_pu_id, "")

        cleaned_data["scope"] = resolved_scope
        cleaned_data["pu_id"] = resolved_pu_id
        if not is_manual_override:
            cleaned_data["pu_name"] = resolved_pu_name
        return cleaned_data

    def save(self, commit=True):
        obj = super().save(commit=False)
        obj.scope = self.cleaned_data.get("scope", obj.scope)
        obj.pu_id = self.cleaned_data.get("pu_id", obj.pu_id)
        obj.pu_name = self.cleaned_data.get("pu_name", obj.pu_name)
        if commit:
            obj.save()
        return obj
