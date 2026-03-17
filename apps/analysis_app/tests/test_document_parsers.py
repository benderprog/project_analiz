from __future__ import annotations

from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZipFile
from types import SimpleNamespace
from unittest.mock import patch

from django.test import SimpleTestCase, override_settings

from apps.analysis_app.document_parsers import (
    PdfTextLayerMissingError,
    UnsupportedDocumentFormatError,
    DocConversionError,
    ExtractedDocument,
    extract_document_text,
)
from apps.analysis_app.services import parse_uploaded_document
from apps.analysis_app.svodka_templates import SegmentAnchors


class DocumentParserTests(SimpleTestCase):
    def test_extract_docx(self):
        from docx import Document

        document = Document()
        document.add_paragraph("Первая строка")
        document.add_paragraph("Вторая строка")

        with NamedTemporaryFile(suffix=".docx") as tmp:
            document.save(tmp.name)
            parsed = extract_document_text(tmp.name)

        self.assertEqual(parsed.source_format, "docx")
        self.assertEqual(parsed.lines, ["Первая строка", "Вторая строка"])
        self.assertTrue(parsed.is_text_based)

    def test_extract_odt(self):
        content_xml = """<?xml version='1.0' encoding='UTF-8'?>
<office:document-content
 xmlns:office='urn:oasis:names:tc:opendocument:xmlns:office:1.0'
 xmlns:text='urn:oasis:names:tc:opendocument:xmlns:text:1.0'>
 <office:body>
  <office:text>
   <text:h>Заголовок</text:h>
   <text:p>Строка ODT</text:p>
  </office:text>
 </office:body>
</office:document-content>"""

        with NamedTemporaryFile(suffix=".odt") as tmp:
            with ZipFile(tmp.name, "w") as archive:
                archive.writestr("content.xml", content_xml)
            parsed = extract_document_text(tmp.name)

        self.assertEqual(parsed.source_format, "odt")
        self.assertEqual(parsed.lines, ["Заголовок", "Строка ODT"])
        self.assertEqual(parsed.text_blocks, ["Заголовок", "Строка ODT"])

    def test_extract_rtf(self):
        rtf_content = r"{\rtf1\ansi\deff0 {\fonttbl {\f0 Times;}}\f0\fs24 Первая строка\par Вторая строка}"

        with NamedTemporaryFile(suffix=".rtf", mode="w", encoding="utf-8") as tmp:
            tmp.write(rtf_content)
            tmp.flush()
            parsed = extract_document_text(tmp.name)

        self.assertEqual(parsed.source_format, "rtf")
        self.assertEqual(parsed.lines, ["Первая строка", "Вторая строка"])
        self.assertEqual(parsed.text_blocks, ["Первая строка", "Вторая строка"])

    def test_extract_pdf_with_text_layer(self):
        pdf_like = b"stream\nBT /F1 12 Tf 72 712 Td (Hello PDF) Tj ET\nendstream"

        with NamedTemporaryFile(suffix=".pdf") as tmp:
            Path(tmp.name).write_bytes(pdf_like)
            parsed = extract_document_text(tmp.name)

        self.assertEqual(parsed.source_format, "pdf")
        self.assertEqual(parsed.lines, ["Hello PDF"])
        self.assertEqual(parsed.text_blocks, ["Hello PDF"])

    def test_extract_pdf_without_text_layer_raises(self):
        with NamedTemporaryFile(suffix=".pdf") as tmp:
            Path(tmp.name).write_bytes(b"stream\nq 10 0 0 10 0 0 cm /Im0 Do Q\nendstream")
            with self.assertRaises(PdfTextLayerMissingError) as error:
                extract_document_text(tmp.name)

        self.assertIn("Scanned/image-only PDF is not supported yet", str(error.exception))

    @override_settings(MIN_EVENT_PARAGRAPH_CHARS=0)
    def test_parse_uploaded_document_uses_structured_blocks_for_rtf(self):
        rtf_content = r"{\rtf1\ansi Универсальный extractor\par строка 2}"

        with NamedTemporaryFile(suffix=".rtf", mode="w", encoding="utf-8") as tmp:
            tmp.write(rtf_content)
            tmp.flush()
            events, meta = parse_uploaded_document(tmp.name, filename="sample.rtf", return_slicing_meta=True)

        self.assertEqual(len(events), 2)
        self.assertEqual(events[0].kind, "paragraph")
        self.assertEqual(events[0].joined_text, "Универсальный extractor")
        self.assertEqual(events[1].joined_text, "строка 2")
        self.assertEqual(meta.get("source_format"), "rtf")
        self.assertEqual(meta.get("method"), "document_text_blocks")

    def test_unsupported_format_raises(self):
        with NamedTemporaryFile(suffix=".txt") as tmp:
            Path(tmp.name).write_text("text", encoding="utf-8")
            with self.assertRaises(UnsupportedDocumentFormatError):
                extract_document_text(tmp.name)

    @override_settings(MIN_EVENT_PARAGRAPH_CHARS=0)
    def test_parse_uploaded_document_keeps_docx_path_regression(self):
        from docx import Document

        document = Document()
        document.add_paragraph("Событие 1")
        document.add_paragraph("Событие 2")

        with NamedTemporaryFile(suffix=".docx") as tmp:
            document.save(tmp.name)
            events, _ = parse_uploaded_document(tmp.name, filename="sample.docx", return_slicing_meta=True)

        self.assertEqual(len(events), 2)
        self.assertEqual([e.joined_text for e in events], ["Событие 1", "Событие 2"])

    @override_settings(MIN_EVENT_PARAGRAPH_CHARS=0)
    def test_parse_uploaded_document_uses_structured_blocks_for_odt(self):
        content_xml = """<?xml version='1.0' encoding='UTF-8'?>
<office:document-content
 xmlns:office='urn:oasis:names:tc:opendocument:xmlns:office:1.0'
 xmlns:text='urn:oasis:names:tc:opendocument:xmlns:text:1.0'>
 <office:body>
  <office:text>
   <text:p>Первое событие</text:p>
   <text:p>Второе событие</text:p>
  </office:text>
 </office:body>
</office:document-content>"""

        with NamedTemporaryFile(suffix=".odt") as tmp:
            with ZipFile(tmp.name, "w") as archive:
                archive.writestr("content.xml", content_xml)
            events, _ = parse_uploaded_document(tmp.name, filename="sample.odt", return_slicing_meta=True)

        self.assertEqual(len(events), 2)
        self.assertEqual([e.joined_text for e in events], ["Первое событие", "Второе событие"])



    @override_settings(MIN_EVENT_PARAGRAPH_CHARS=0, TEMPLATE_MIN_SLICE_CHARS=1, SKIP_SEMANTIC_MODEL=True)
    def test_parse_uploaded_document_applies_docx_template_for_odt_summary(self):
        from docx import Document

        content_xml = """<?xml version='1.0' encoding='UTF-8'?>
<office:document-content
 xmlns:office='urn:oasis:names:tc:opendocument:xmlns:office:1.0'
 xmlns:text='urn:oasis:names:tc:opendocument:xmlns:text:1.0'>
 <office:body>
  <office:text>
   <text:p>start anchor</text:p>
   <text:p>Внутри шаблона ODT</text:p>
   <text:p>end anchor</text:p>
  </office:text>
 </office:body>
</office:document-content>"""

        template = Document()
        template.add_paragraph("header")
        template.add_paragraph("[BEGIN]")
        template.add_paragraph("body")
        template.add_paragraph("[END]")
        template.add_paragraph("footer")

        with NamedTemporaryFile(suffix=".odt") as summary_tmp, NamedTemporaryFile(suffix=".docx") as template_tmp:
            with ZipFile(summary_tmp.name, "w") as archive:
                archive.writestr("content.xml", content_xml)
            template.save(template_tmp.name)

            fake_template = SimpleNamespace(
                file=SimpleNamespace(path=template_tmp.name, name="template.docx"),
                begin_marker="[BEGIN]",
                end_marker="[END]",
                anchor_match_threshold=0.6,
                scope="pu",
                pu_id="1",
                template_id="fake-template",
            )
            with patch("apps.analysis_app.svodka_templates.get_template_for_run", return_value=fake_template), patch(
                "apps.analysis_app.svodka_templates.build_template_segments",
                return_value=[SegmentAnchors(start_anchor_text="start anchor", end_anchor_text="end anchor", index=0)],
            ):
                events, meta = parse_uploaded_document(
                    summary_tmp.name,
                    filename="summary.odt",
                    selected_pu_id="1",
                    selected_pu_name="ПУ",
                    return_slicing_meta=True,
                )

        self.assertEqual([e.joined_text for e in events], ["Внутри шаблона ODT"])
        self.assertEqual(meta.get("slice_strategy"), "template_anchors")
        self.assertEqual(meta.get("anchors_matched"), 1)

    @override_settings(MIN_EVENT_PARAGRAPH_CHARS=0, TEMPLATE_MIN_SLICE_CHARS=1, SKIP_SEMANTIC_MODEL=True)
    def test_parse_uploaded_document_applies_docx_template_for_docx_summary_regression(self):
        from docx import Document

        summary = Document()
        summary.add_paragraph("start anchor")
        summary.add_paragraph("inside docx")
        summary.add_paragraph("end anchor")

        template = Document()
        template.add_paragraph("header")
        template.add_paragraph("[BEGIN]")
        template.add_paragraph("body")
        template.add_paragraph("[END]")
        template.add_paragraph("footer")

        with NamedTemporaryFile(suffix=".docx") as summary_tmp, NamedTemporaryFile(suffix=".docx") as template_tmp:
            summary.save(summary_tmp.name)
            template.save(template_tmp.name)

            fake_template = SimpleNamespace(
                file=SimpleNamespace(path=template_tmp.name, name="template.docx"),
                begin_marker="[BEGIN]",
                end_marker="[END]",
                anchor_match_threshold=0.6,
                scope="pu",
                pu_id="1",
                template_id="fake-template",
            )
            with patch("apps.analysis_app.svodka_templates.get_template_for_run", return_value=fake_template), patch(
                "apps.analysis_app.svodka_templates.build_template_segments",
                return_value=[SegmentAnchors(start_anchor_text="start anchor", end_anchor_text="end anchor", index=0)],
            ):
                events, meta = parse_uploaded_document(
                    summary_tmp.name,
                    filename="summary.docx",
                    selected_pu_id="1",
                    selected_pu_name="ПУ",
                    return_slicing_meta=True,
                )

        self.assertEqual([e.joined_text for e in events], ["inside docx"])
        self.assertEqual(meta.get("slice_strategy"), "template_anchors")
        self.assertEqual(meta.get("anchors_matched"), 1)




    def test_extract_doc_raises_controlled_error_when_converter_missing(self):
        with NamedTemporaryFile(suffix=".doc") as tmp, patch("apps.analysis_app.document_parsers.shutil.which", return_value=None):
            Path(tmp.name).write_bytes(b"legacy doc payload")
            with self.assertRaises(DocConversionError) as error:
                extract_document_text(tmp.name)

        self.assertIn("local converter 'soffice'", str(error.exception))

    def test_extract_doc_uses_headless_local_conversion(self):
        from docx import Document

        with NamedTemporaryFile(suffix=".doc") as tmp:
            Path(tmp.name).write_bytes(b"legacy doc payload")

            def _fake_convert(cmd, capture_output, text, check):
                self.assertIn("--headless", cmd)
                input_path = Path(cmd[-1])
                out_dir = Path(cmd[cmd.index("--outdir") + 1])
                output_path = out_dir / f"{input_path.stem}.docx"
                converted = Document()
                converted.add_paragraph("DOC line 1")
                converted.add_paragraph("DOC line 2")
                converted.save(output_path)
                return SimpleNamespace(returncode=0, stdout="ok", stderr="")

            with patch("apps.analysis_app.document_parsers.shutil.which", return_value="/usr/bin/soffice"), patch(
                "apps.analysis_app.document_parsers.subprocess.run",
                side_effect=_fake_convert,
            ):
                parsed = extract_document_text(tmp.name)

        self.assertEqual(parsed.source_format, "doc")
        self.assertEqual(parsed.lines, ["DOC line 1", "DOC line 2"])
        self.assertEqual(parsed.meta.get("converted_from"), "doc")

    def test_extract_doc_falls_back_to_single_docx_in_output_dir(self):
        from docx import Document

        with NamedTemporaryFile(suffix=".DOC") as tmp:
            Path(tmp.name).write_bytes(b"legacy doc payload")

            def _fake_convert(cmd, capture_output, text, check):
                out_dir = Path(cmd[cmd.index("--outdir") + 1])
                converted = Document()
                converted.add_paragraph("Converted fallback")
                converted.save(out_dir / "unexpected output name.docx")
                return SimpleNamespace(returncode=0, stdout="ok", stderr="")

            with patch("apps.analysis_app.document_parsers.shutil.which", return_value="/usr/bin/soffice"), patch(
                "apps.analysis_app.document_parsers.subprocess.run",
                side_effect=_fake_convert,
            ):
                parsed = extract_document_text(tmp.name)

        self.assertEqual(parsed.lines, ["Converted fallback"])
        self.assertEqual(parsed.meta.get("conversion_output_file"), "unexpected output name.docx")

    def test_extract_doc_includes_conversion_diagnostics_on_missing_output(self):
        with NamedTemporaryFile(suffix=".doc") as tmp:
            Path(tmp.name).write_bytes(b"legacy doc payload")

            with patch("apps.analysis_app.document_parsers.shutil.which", return_value="/usr/bin/soffice"), patch(
                "apps.analysis_app.document_parsers.subprocess.run",
                return_value=SimpleNamespace(returncode=0, stdout="converter stdout", stderr="converter stderr"),
            ):
                with self.assertRaises(DocConversionError) as error:
                    extract_document_text(tmp.name)

        message = str(error.exception)
        self.assertIn("command:", message)
        self.assertIn("returncode:", message)
        self.assertIn("stdout:", message)
        self.assertIn("stderr:", message)
        self.assertIn("output_dir_listing:", message)

    @override_settings(MIN_EVENT_PARAGRAPH_CHARS=0)
    def test_parse_uploaded_document_supports_doc_summary_via_conversion(self):
        with NamedTemporaryFile(suffix=".doc") as tmp:
            Path(tmp.name).write_bytes(b"legacy doc payload")
            with patch(
                "apps.analysis_app.document_parsers._extract_doc",
                return_value=ExtractedDocument(
                    source_format="doc",
                    text="DOC event 1\nDOC event 2",
                    lines=["DOC event 1", "DOC event 2"],
                    text_blocks=["DOC event 1", "DOC event 2"],
                    meta={"converted_from": "doc"},
                ),
            ):
                events, meta = parse_uploaded_document(tmp.name, filename="summary.doc", return_slicing_meta=True)

        self.assertEqual([e.joined_text for e in events], ["DOC event 1", "DOC event 2"])
        self.assertEqual(meta.get("source_format"), "doc")

    @override_settings(MIN_EVENT_PARAGRAPH_CHARS=0, TEMPLATE_MIN_SLICE_CHARS=1, SKIP_SEMANTIC_MODEL=True)
    def test_parse_uploaded_document_supports_doc_template_via_conversion(self):
        from docx import Document

        summary = Document()
        summary.add_paragraph("anchor start")
        summary.add_paragraph("inside docx")
        summary.add_paragraph("anchor end")

        with NamedTemporaryFile(suffix=".docx") as summary_tmp, NamedTemporaryFile(suffix=".doc") as template_tmp:
            summary.save(summary_tmp.name)
            Path(template_tmp.name).write_bytes(b"legacy template payload")

            fake_template = SimpleNamespace(
                file=SimpleNamespace(path=template_tmp.name, name="template.doc"),
                begin_marker="[BEGIN]",
                end_marker="[END]",
                anchor_match_threshold=0.6,
                scope="pu",
                pu_id="1",
                template_id="fake-template",
            )

            with patch("apps.analysis_app.svodka_templates.get_template_for_run", return_value=fake_template), patch(
                "apps.analysis_app.document_parsers._extract_doc",
                return_value=ExtractedDocument(
                    source_format="doc",
                    text="anchor start\n[BEGIN]\nbody\n[END]\nanchor end",
                    lines=["anchor start", "[BEGIN]", "body", "[END]", "anchor end"],
                    text_blocks=["anchor start", "[BEGIN]", "body", "[END]", "anchor end"],
                    meta={"converted_from": "doc"},
                ),
            ):
                events, meta = parse_uploaded_document(
                    summary_tmp.name,
                    filename="summary.docx",
                    selected_pu_id="1",
                    selected_pu_name="ПУ",
                    return_slicing_meta=True,
                )

        self.assertEqual([e.joined_text for e in events], ["inside docx"])
        self.assertEqual(meta.get("slice_strategy"), "template_anchors")


    @override_settings(MIN_EVENT_PARAGRAPH_CHARS=0)
    def test_parse_uploaded_document_uses_blocks_for_pdf(self):
        pdf_like = b"stream\nBT /F1 12 Tf 72 712 Td (PDF block 1) Tj ET\nendstream\nstream\nBT /F1 12 Tf 72 680 Td (PDF block 2) Tj ET\nendstream"

        with NamedTemporaryFile(suffix=".pdf") as tmp:
            Path(tmp.name).write_bytes(pdf_like)
            events, _ = parse_uploaded_document(tmp.name, filename="sample.pdf", return_slicing_meta=True)

        self.assertEqual(len(events), 2)
        self.assertEqual([e.joined_text for e in events], ["PDF block 1", "PDF block 2"])
