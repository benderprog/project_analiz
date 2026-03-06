from tempfile import NamedTemporaryFile
from types import SimpleNamespace
import os
from unittest.mock import patch

from django.test import SimpleTestCase, override_settings
from docx import Document

from apps.analysis_app.svodka_templates import (
    DocElement,
    SegmentAnchors,
    apply_template_segments,
    build_template_segments,
    parse_template_marker_blocks,
    slice_by_markers,
    normalize_anchor_text,
    slice_document_for_run,
)




class AnchorNormalizationTests(SimpleTestCase):
    def test_normalize_anchor_text_handles_case_punctuation_and_yo(self):
        self.assertEqual(normalize_anchor_text("Ёлка, №1"), "елка 1")
        self.assertEqual(normalize_anchor_text("елка 1"), "елка 1")

class MarkerSlicingTests(SimpleTestCase):
    def test_two_segments_with_gap_keeps_only_segment_content(self):
        elements = [
            DocElement(kind="paragraph", text="[BEGIN]"),
            DocElement(kind="paragraph", text="A1"),
            DocElement(kind="paragraph", text="[END]"),
            DocElement(kind="paragraph", text="GAP"),
            DocElement(kind="paragraph", text="[BEGIN]"),
            DocElement(kind="paragraph", text="B1"),
            DocElement(kind="paragraph", text="[END]"),
        ]

        kept, debug = slice_by_markers(elements, "[BEGIN]", "[END]")

        self.assertEqual([item.text for item in kept], ["A1", "B1"])
        self.assertTrue(debug["markers_found"])
        self.assertEqual(debug["segments_count"], 2)
        self.assertEqual(debug["kept_elements"], 2)

    def test_end_without_begin_keeps_nothing_for_direct_marker_slicing(self):
        elements = [
            DocElement(kind="paragraph", text="X"),
            DocElement(kind="paragraph", text="[END]"),
            DocElement(kind="paragraph", text="Y"),
        ]

        kept, debug = slice_by_markers(elements, "[BEGIN]", "[END]")

        self.assertEqual(kept, [])
        self.assertTrue(debug["markers_found"])


class TemplateBlockParsingTests(SimpleTestCase):
    def test_single_segment_extracts_pre_and_post_anchors(self):
        blocks, warnings = parse_template_marker_blocks("pre\n[BEGIN]\nbody\n[END]\npost")
        self.assertEqual(len(blocks), 1)
        self.assertEqual(blocks[0].pre_anchor_text, "pre")
        self.assertEqual(blocks[0].post_anchor_text, "post")
        self.assertEqual(warnings, [])

    def test_multiple_segments_supported(self):
        text = "a\n[BEGIN]\nA\n[END]\nb\n[BEGIN]\nB\n[END]\nc"
        blocks, _ = parse_template_marker_blocks(text)
        self.assertEqual(len(blocks), 2)

    def test_begin_without_end_open_ended(self):
        blocks, warnings = parse_template_marker_blocks("pre\n[BEGIN]\nbody")
        self.assertEqual(len(blocks), 1)
        self.assertIsNone(blocks[0].end_marker_span)
        self.assertTrue(any("BEGIN without END" in item for item in warnings))

    def test_end_without_begin_is_ignored(self):
        blocks, warnings = parse_template_marker_blocks("x\n[END]\ny")
        self.assertEqual(blocks, [])
        self.assertTrue(any("END without BEGIN" in item for item in warnings))


class SvodkaTemplateSlicingTests(SimpleTestCase):
    @override_settings(SKIP_SEMANTIC_MODEL=True)
    def test_exact_anchors_in_report_slice_between_lines(self):
        anchors = [SegmentAnchors(start_anchor_text="A start", end_anchor_text="A end", index=0)]
        target = [
            DocElement(kind="paragraph", text="A start"),
            DocElement(kind="paragraph", text="inside"),
            DocElement(kind="paragraph", text="A end"),
        ]

        sliced, applied, failed, _, _ = apply_template_segments(target, anchors)

        self.assertEqual(applied, 1)
        self.assertEqual(failed, 0)
        self.assertEqual([e.text for e in sliced], ["inside"])

    @override_settings(SKIP_SEMANTIC_MODEL=True)
    def test_open_ended_segment_slices_to_end(self):
        anchors = [SegmentAnchors(start_anchor_text="A start", end_anchor_text=None, is_open_ended=True, index=0)]
        target = [
            DocElement(kind="paragraph", text="A start"),
            DocElement(kind="paragraph", text="inside"),
            DocElement(kind="paragraph", text="tail"),
        ]

        sliced, applied, failed, _, _ = apply_template_segments(target, anchors)
        self.assertEqual(applied, 1)
        self.assertEqual(failed, 0)
        self.assertEqual([e.text for e in sliced], ["inside", "tail"])

    @override_settings(SKIP_SEMANTIC_MODEL=True)
    def test_per_template_threshold_affects_match_acceptance(self):
        anchors = [
            SegmentAnchors(
                start_anchor_text="дата нарушения место служба",
                end_anchor_text=None,
                is_open_ended=True,
                index=0,
            )
        ]
        target = [
            DocElement(kind="paragraph", text="Дата нарушения и место"),
            DocElement(kind="paragraph", text="тело"),
            DocElement(kind="paragraph", text="хвост"),
        ]

        sliced_high, applied_high, failed_high, _, _ = apply_template_segments(target, anchors, base_threshold=0.65)
        sliced_low, applied_low, failed_low, _, _ = apply_template_segments(target, anchors, base_threshold=0.6)

        self.assertEqual(sliced_high, [])
        self.assertEqual(applied_high, 0)
        self.assertEqual(failed_high, 1)

        self.assertEqual(applied_low, 1)
        self.assertEqual(failed_low, 0)
        self.assertEqual([e.text for e in sliced_low], ["тело", "хвост"])

    @override_settings(SKIP_SEMANTIC_MODEL=True)
    def test_paraphrased_anchor_lexical_match(self):
        anchors = [SegmentAnchors(start_anchor_text="дата нарушения", end_anchor_text="итоги проверки", index=0)]
        target = [
            DocElement(kind="paragraph", text="Дата нарушения и место"),
            DocElement(kind="paragraph", text="тело"),
            DocElement(kind="paragraph", text="Итоги проверки по делу"),
        ]

        sliced, applied, failed, _, _ = apply_template_segments(target, anchors)
        self.assertEqual(applied, 1)
        self.assertEqual(failed, 0)
        self.assertEqual([e.text for e in sliced], ["тело"])

    def test_build_template_segments_supports_open_ended(self):
        template_elements = [
            DocElement(kind="paragraph", text="start line"),
            DocElement(kind="paragraph", text="[BEGIN]"),
            DocElement(kind="paragraph", text="body"),
        ]
        segments = build_template_segments(template_elements, "[BEGIN]", "[END]")
        self.assertEqual(len(segments), 1)
        self.assertTrue(segments[0].is_open_ended)

    def test_regression_literal_markers_path_unchanged(self):
        elements = [
            DocElement(kind="paragraph", text="x"),
            DocElement(kind="paragraph", text="[BEGIN]"),
            DocElement(kind="paragraph", text="y"),
            DocElement(kind="paragraph", text="[END]"),
            DocElement(kind="paragraph", text="z"),
        ]
        kept, _ = slice_by_markers(elements, "[BEGIN]", "[END]")
        self.assertEqual([item.text for item in kept], ["y"])


class SliceDocumentForRunFallbackTests(SimpleTestCase):
    def _write_docx(self, paragraphs: list[str]) -> str:
        document = Document()
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            document.save(tmp.name)
            return tmp.name

    @override_settings(SKIP_SEMANTIC_MODEL=True, MIN_EVENT_PARAGRAPH_CHARS=20)
    def test_anchor_miss_marks_meta_and_falls_back_to_full_document(self):
        report_path = self._write_docx(["очень длинный абзац " * 8, "второй длинный абзац " * 8])
        template_path = self._write_docx(["pre", "[BEGIN]", "body", "[END]", "post"])
        try:
            report_doc = Document(report_path)
            fake_template = SimpleNamespace(
                file=SimpleNamespace(path=template_path),
                begin_marker="[BEGIN]",
                end_marker="[END]",
                anchor_match_threshold=0.6,
            )
            with patch("apps.analysis_app.svodka_templates.get_template_for_run", return_value=fake_template), patch(
                "apps.analysis_app.svodka_templates.build_template_segments",
                return_value=[SegmentAnchors(start_anchor_text="missing", end_anchor_text="missing-end", index=0)],
            ):
                kept, info = slice_document_for_run(report_doc, selected_pu_id="1", selected_pu_name="ПУ")

            self.assertEqual(info.slicing_strategy, "none")
            self.assertTrue(info.anchors_missing)
            self.assertEqual(info.anchors_expected, 1)
            self.assertEqual(info.anchors_matched, 0)
            self.assertEqual(len(kept), info.total_elements)
        finally:
            os.unlink(report_path)
            os.unlink(template_path)

    @override_settings(SKIP_SEMANTIC_MODEL=True)
    def test_when_anchors_match_method_is_template_anchors(self):
        report_path = self._write_docx(["start anchor", "inside segment text", "end anchor"])
        template_path = self._write_docx(["pre", "[BEGIN]", "body", "[END]", "post"])
        try:
            report_doc = Document(report_path)
            fake_template = SimpleNamespace(
                file=SimpleNamespace(path=template_path),
                begin_marker="[BEGIN]",
                end_marker="[END]",
                anchor_match_threshold=0.6,
            )
            with patch("apps.analysis_app.svodka_templates.get_template_for_run", return_value=fake_template), patch(
                "apps.analysis_app.svodka_templates.build_template_segments",
                return_value=[SegmentAnchors(start_anchor_text="start anchor", end_anchor_text="end anchor", index=0)],
            ):
                kept, info = slice_document_for_run(report_doc, selected_pu_id="1", selected_pu_name="ПУ")

            self.assertEqual(info.slicing_strategy, "template_anchors")
            self.assertEqual(info.anchors_matched, 1)
            self.assertEqual([item.text for item in kept], ["inside segment text"])
        finally:
            os.unlink(report_path)
            os.unlink(template_path)

    def test_report_markers_keep_priority(self):
        report_path = self._write_docx(["prefix", "[BEGIN]", "inside", "[END]", "suffix"])
        template_path = self._write_docx(["pre", "[BEGIN]", "body", "[END]", "post"])
        try:
            report_doc = Document(report_path)
            fake_template = SimpleNamespace(
                file=SimpleNamespace(path=template_path),
                begin_marker="[BEGIN]",
                end_marker="[END]",
                anchor_match_threshold=0.6,
            )
            with patch("apps.analysis_app.svodka_templates.get_template_for_run", return_value=fake_template):
                kept, info = slice_document_for_run(report_doc, selected_pu_id="1", selected_pu_name="ПУ")

            self.assertEqual(info.slicing_strategy, "report_markers")
            self.assertEqual([item.text for item in kept], ["inside"])
        finally:
            os.unlink(report_path)
            os.unlink(template_path)
