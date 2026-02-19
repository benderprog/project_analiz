from __future__ import annotations

from tempfile import NamedTemporaryFile
import os

from django.test import SimpleTestCase, override_settings
from docx import Document

from apps.analysis_app.services import parse_docx


class ParseDocxParagraphFilterTests(SimpleTestCase):
    def _write_docx(self, paragraphs: list[str], table_rows: list[list[str]] | None = None, paragraph_after: str | None = None) -> str:
        document = Document()
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        if table_rows is not None:
            cols = len(table_rows[0]) if table_rows else 1
            table = document.add_table(rows=0, cols=cols)
            for row_cells in table_rows:
                row = table.add_row()
                for idx, value in enumerate(row_cells):
                    row.cells[idx].text = value
        if paragraph_after is not None:
            document.add_paragraph(paragraph_after)

        with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            document.save(tmp.name)
            return tmp.name

    def test_parse_docx_skips_short_blocks(self):
        long_text = "Событие " + ("очень важное " * 12)
        file_path = self._write_docx([
            "Заголовок",
            "   ",
            "---",
            long_text,
        ])

        try:
            with self.assertLogs("apps.analysis_app.services", level="INFO") as captured:
                events = parse_docx(file_path)

            self.assertEqual([event.joined_text for event in events], [long_text.strip()])
            self.assertTrue(any("docx split: total=4, kept=1, skipped_short=3" in msg for msg in captured.output))
        finally:
            os.unlink(file_path)

    @override_settings(MIN_EVENT_PARAGRAPH_CHARS=20)
    def test_parse_docx_threshold_is_configurable(self):
        file_path = self._write_docx([
            "Это абзац длиннее двадцати символов.",
            "Короткий",
        ])

        try:
            events = parse_docx(file_path)

            self.assertEqual([event.joined_text for event in events], ["Это абзац длиннее двадцати символов."])
        finally:
            os.unlink(file_path)

    @override_settings(MIN_EVENT_PARAGRAPH_CHARS=100)
    def test_parse_docx_uses_normalized_length(self):
        noisy_short = "Слишком" + ("\n\n     " * 40) + "короткий"
        normalized_long = ("Длинный факт о событии " * 6).strip()
        file_path = self._write_docx([noisy_short, normalized_long])

        try:
            events = parse_docx(file_path)

            self.assertEqual([event.joined_text for event in events], [normalized_long])
        finally:
            os.unlink(file_path)

    @override_settings(MIN_EVENT_PARAGRAPH_CHARS=100)
    def test_parse_docx_keeps_document_order_for_paragraphs_and_table_rows(self):
        paragraph_a = "Параграф A " + ("длинный текст " * 10)
        paragraph_b = "Параграф B " + ("длинный текст " * 10)
        row_1 = ["ячейка 1", "ячейка 2", "ячейка 3 " + ("данные " * 12)]
        row_2 = ["вторая 1", "вторая 2", "вторая 3 " + ("описание " * 12)]
        file_path = self._write_docx([paragraph_a], [row_1, row_2], paragraph_after=paragraph_b)

        try:
            events = parse_docx(file_path)

            self.assertEqual([event.kind for event in events], ["paragraph", "table_row", "table_row", "paragraph"])
            self.assertEqual(events[1].cells, ["ячейка 1", "ячейка 2", row_1[2].strip()])
            self.assertEqual(events[1].joined_text, f"ячейка 1 ячейка 2 {row_1[2].strip()}")
            self.assertEqual(events[2].joined_text, f"вторая 1 вторая 2 {row_2[2].strip()}")
            self.assertEqual(events[0].joined_text, paragraph_a.strip())
            self.assertEqual(events[3].joined_text, paragraph_b.strip())
        finally:
            os.unlink(file_path)

    @override_settings(MIN_EVENT_PARAGRAPH_CHARS=100)
    def test_parse_docx_skips_short_and_empty_table_rows(self):
        long_row = ["Колонка 1 " + ("длинно " * 8), "Колонка 2 " + ("длинно " * 8)]
        file_path = self._write_docx(
            ["Вводный абзац " + ("длинный " * 12)],
            [
                ["", "   "],
                ["коротко", "мало"],
                long_row,
            ],
        )

        try:
            events = parse_docx(file_path)

            self.assertEqual(len(events), 2)
            self.assertEqual(events[1].kind, "table_row")
            self.assertEqual(events[1].joined_text, f"{long_row[0].strip()} {long_row[1].strip()}")
        finally:
            os.unlink(file_path)


    @override_settings(MIN_EVENT_PARAGRAPH_CHARS=60)
    def test_parse_docx_detects_table_header_and_skips_it_as_event(self):
        header = ["Дата", "Время", "Событие"]
        row_1 = ["01.01.2026", "12:00", "Описание события " + ("очень важное " * 6)]
        row_2 = ["02.01.2026", "13:00", "Второе описание события " + ("детали " * 8)]
        file_path = self._write_docx([], [header, row_1, row_2])

        try:
            events = parse_docx(file_path)

            self.assertEqual(len(events), 2)
            self.assertEqual([event.kind for event in events], ["table_row", "table_row"])
            self.assertEqual(events[0].cells, row_1)
            self.assertEqual(events[1].cells, row_2)
            self.assertEqual(events[0].table_header_cells, header)
            self.assertEqual(events[1].table_header_cells, header)
        finally:
            os.unlink(file_path)

    @override_settings(MIN_EVENT_PARAGRAPH_CHARS=60)
    def test_parse_docx_table_without_header_keeps_first_row_as_event(self):
        row_1 = [
            "Первое длинное описание события " + ("данные " * 8),
            "Продолжение первого описания " + ("детали " * 7),
        ]
        row_2 = [
            "Второе длинное описание события " + ("данные " * 8),
            "Продолжение второго описания " + ("детали " * 7),
        ]
        file_path = self._write_docx([], [row_1, row_2])

        try:
            events = parse_docx(file_path)

            self.assertEqual(len(events), 2)
            self.assertEqual(events[0].cells, row_1)
            self.assertIsNone(events[0].table_header_cells)
            self.assertIsNone(events[1].table_header_cells)
        finally:
            os.unlink(file_path)
